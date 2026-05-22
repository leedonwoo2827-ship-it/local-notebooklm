"""OCR 없는 가벼운 텍스트 추출 — PDF(텍스트 추출 가능본), Docx 만 지원.

스캔/이미지 PDF 는 본 앱이 지원하지 않는다 (사용자 결정: OCR 의존 경로 제거).
빈 결과를 돌려주면 `sources_panel` 이 사용자에게 외부 OCR 사용을 안내한다.
"""
from __future__ import annotations

from pathlib import Path


def pdf_to_text(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        if t.strip():
            parts.append(t)
    return "\n\n".join(parts).strip()


def docx_to_text(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)
    return "\n".join(parts).strip()
