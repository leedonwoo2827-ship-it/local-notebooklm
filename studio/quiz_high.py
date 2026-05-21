"""고난도 30문제 — 책 단원별 균등 분배 + 분할 호출 + xlsx 저장.

PDF 파일명 규칙: `<순번>_<단원명>_<시작쪽>-<끝쪽>.pdf`
예: `01_채소_12-134.pdf`, `02_과수_135-262.pdf`
해당 패턴 매칭 파일이 ≥ 2 개면 단원별로 분할 호출, 아니면 균등 5등분 폴백.

이 산출물은 `hidden=True` 라 기본적으로 숨김. `.env` 에
`STUDIO_VISIBLE=quiz_high` 추가해야 본인 환경에서만 버튼이 노출된다.
"""
from __future__ import annotations

import json
import re
import time
from pathlib import Path

from core.rag import NotebookPaths

from ._base import ArtifactMeta, ArtifactResult, load_prompt
from ._xlsx import write_table_xlsx
from .quiz import _strip_fence

META = ArtifactMeta(
    key="quiz_high",
    title="고난도 30문제",
    icon="🎯",
    order=55,
    model_profile="strong",
    description="책 단원별 균등 분배(분할 호출). 본인 환경에서만 노출(hidden).",
    hidden=True,
)

TOTAL_QUESTIONS = 30
FALLBACK_SECTIONS = 5  # 단원 정보 없을 때 균등 N등분

_SECTION_PATTERN = re.compile(r"^\d+_(.+?)_(\d+)-(\d+)\.(?:pdf|PDF)$")


def _discover_sections(notebook_name: str) -> list[dict]:
    """노트북 sources 의 PDF 파일명에서 단원 정보 추출.

    매칭되는 파일이 2개 이상일 때만 단원 목록 반환. 그 미만이면 빈 리스트
    (호출 측에서 균등 N등분 폴백으로 분기).
    """
    paths = NotebookPaths.for_notebook(notebook_name)
    if not paths.sources.exists():
        return []
    sections: list[dict] = []
    for p in sorted(paths.sources.iterdir()):
        m = _SECTION_PATTERN.match(p.name)
        if m:
            sections.append({
                "name": m.group(1),
                "start": int(m.group(2)),
                "end": int(m.group(3)),
                "file": p.name,
            })
    return sections if len(sections) >= 2 else []


def _section_spec(
    *,
    section_name: str,
    section_count: int,
    total: int,
    index: int,
    total_sections: int,
    prior_titles: list[str],
    page_range: str | None = None,
) -> str:
    prior_block = (
        "이미 다른 단원에서 출제된 문항 주제(중복 출제 금지):\n"
        + "\n".join(f"  - {t}" for t in prior_titles)
        if prior_titles else
        "이전 단원에서 출제된 문항 없음 (이번이 첫 단원)."
    )
    page_line = f" (책 {page_range}쪽)" if page_range else ""
    return (
        "[단원별 분할 출제 모드]\n"
        f"전체 {total}문항을 {total_sections}단원에 균등 분배 중. "
        f"이번 호출: **{index}/{total_sections} — 「{section_name}」**{page_line} "
        f"영역에서만 {section_count}문항.\n"
        f"이 단원과 무관한 내용은 출제하지 마라. RAG 자료 중 「{section_name}」 "
        "관련 청크 위주로 참고하라.\n"
        f"{prior_block}\n"
        f"출력 JSON 의 `questions` 배열에는 **이번 단원 {section_count}문항만** 담아라. "
        f"각 문항의 `source` 필드에는 「{section_name}」을 명시하라."
    )


async def _call_section(rag, instruction: str) -> list[dict]:
    raw = await rag.aquery(instruction, mode="hybrid", top_k=30)
    if not raw:
        return []
    try:
        bundle = json.loads(_strip_fence(raw))
    except json.JSONDecodeError:
        return []
    if isinstance(bundle, dict):
        return bundle.get("questions", []) or []
    if isinstance(bundle, list):
        return bundle
    return []


def _pad4(items: list) -> list:
    """보기/출처가 4개 미만이면 빈 문자열로 패딩."""
    return list(items) + [""] * (4 - len(items))


def _markdown(questions: list[dict]) -> str:
    out: list[str] = []
    for i, q in enumerate(questions, 1):
        out.append(f"### Q{i}. {q.get('question', '')}")
        choices = _pad4(q.get("choices", []))
        for idx in range(4):
            out.append(f"{idx + 1}. {choices[idx]}")
        out.append("")
        srcs = _pad4(q.get("choice_sources", []))
        check_block = _check_row(srcs).replace("\n", "  \n")  # md 강제 줄바꿈
        out.append(
            "<details><summary>정답·해설·출처·검수</summary>\n\n"
            f"**정답**: {q.get('answer', '?')}\n\n"
            f"**해설**: {q.get('rationale', '')}\n\n"
            f"**출처**: {q.get('source', '')}\n\n"
            f"**검수**:  \n{check_block}\n\n"
            "</details>"
        )
        out.append("")
    return "\n".join(out).strip()


def _check_row(srcs: list[str]) -> str:
    """보기별 좌표를 줄바꿈으로 나열한다.
    예:
        ① [p.78]
        ② [p.78 (결구 불량 관련)]
        ③ [p.78 (착과 불량 선택지 변형)]
        ④ [p.78 (종자 발아 불량 선택지 변형)]
    """
    marks = "①②③④"
    return "\n".join(f"{marks[k]} [{srcs[k]}]" for k in range(4) if srcs[k])


def _xlsx_rows(questions: list[dict]) -> list[list]:
    """xlsx 행. 보기 셀에는 본문만. 검수 컬럼에 선지별 좌표."""
    rows: list[list] = []
    for i, q in enumerate(questions, 1):
        choices = _pad4(q.get("choices", []))
        srcs = _pad4(q.get("choice_sources", []))
        rows.append([
            i,
            q.get("question", ""),
            choices[0], choices[1], choices[2], choices[3],
            q.get("answer", ""),
            q.get("rationale", ""),
            q.get("source", ""),
            _check_row(srcs),
        ])
    return rows


def _write_docx(path: Path, questions: list[dict], deck_title: str) -> Path:
    """문항 1개당 표 1개로 docx 출력. 검수자 친화적 카드 형식.

    표 구조 (10행 2열):
      행1:  Q01    [종합 출처]            (병합 + bold 헤더)
      행2:  문제   | (질문 본문)
      행3~6: ①~④  | (보기 본문만, 좌표 없음)
      행7:  정답   | 1~4
      행8:  해설   | (보기별 거짓 부분 + 함정 패턴)
      행9:  출처   | (단원명·페이지 종합)
      행10: 검수   | ① p.78  ② p.79  ③ p.80  ④ p.78 (선지별 좌표 모아두기)
    """
    from docx import Document
    from docx.shared import Pt

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(deck_title, level=1)
    sub = doc.add_paragraph()
    sub.add_run(
        f"총 {len(questions)}문항 · 4지선다 + 정답 + 해설 + 출처"
    ).italic = True
    doc.add_paragraph()

    for i, q in enumerate(questions, 1):
        choices = _pad4(q.get("choices", []))
        srcs = _pad4(q.get("choice_sources", []))
        table = doc.add_table(rows=10, cols=2)
        try:
            table.style = "Light Grid Accent 1"
        except (KeyError, AttributeError):
            pass  # 스타일 없으면 기본

        # 헤더 행 (병합)
        hdr_left, hdr_right = table.rows[0].cells
        hdr = hdr_left.merge(hdr_right)
        hdr.text = ""
        run = hdr.paragraphs[0].add_run(
            f"Q{i:02d}    {q.get('source', '')}"
        )
        run.bold = True
        run.font.size = Pt(12)

        body = [
            ("문제", q.get("question", "")),
            ("①", choices[0]),
            ("②", choices[1]),
            ("③", choices[2]),
            ("④", choices[3]),
            ("정답", str(q.get("answer", ""))),
            ("해설", q.get("rationale", "")),
            ("출처", q.get("source", "")),
            ("검수", _check_row(srcs)),
        ]
        for ri, (label, content) in enumerate(body, 1):
            c0 = table.rows[ri].cells[0]
            c1 = table.rows[ri].cells[1]
            c0.text = label
            # 셀 내 줄바꿈은 paragraph 분리로 처리 (cell.text 에 \n 넣으면 한 줄로 표시됨)
            lines = (content or "").split("\n") if content else [""]
            c1.text = lines[0]
            for line in lines[1:]:
                c1.add_paragraph(line)
            for run in c0.paragraphs[0].runs:
                run.bold = True

        doc.add_paragraph()  # 문항 사이 빈 줄

    doc.save(path)
    return path


async def generate(rag, context: dict) -> ArtifactResult:
    notebook_name = context.get("notebook_name", "default")
    total = int(context.get("count", TOTAL_QUESTIONS))

    sections = _discover_sections(notebook_name)
    prompt_tpl = load_prompt(META.key) or (
        f"노트북 핵심에서 4지선다 문항 {{{{N}}}}개를 출제하라. "
        "각 항목 키: question, choices(4개), answer(1~4), rationale, source. "
        "최상위는 {\"questions\": [...]} JSON. JSON 외 텍스트 금지.\n\n{{SECTION_SPEC}}"
    )

    all_questions: list[dict] = []
    prior_titles: list[str] = []

    if sections:
        n = len(sections)
        per = total // n
        remainder = total - per * n  # 앞 단원에 +1 씩 분배
        print(f"[QuizHigh] 단원 분할 모드: {n}개 단원 × {per}문항 (+{remainder} 보정)",
              flush=True)
        for i, sec in enumerate(sections, 1):
            count = per + (1 if i <= remainder else 0)
            spec = _section_spec(
                section_name=sec["name"],
                section_count=count,
                total=total,
                index=i,
                total_sections=n,
                prior_titles=prior_titles,
                page_range=f"{sec['start']}-{sec['end']}",
            )
            instruction = (
                prompt_tpl.replace("{{N}}", str(count))
                .replace("{{SECTION_SPEC}}", spec)
            )
            print(f"[QuizHigh] 단원 {i}/{n} 「{sec['name']}」 시작 ({count}문항)",
                  flush=True)
            qs = await _call_section(rag, instruction)
            print(f"[QuizHigh] 단원 {i}/{n} 「{sec['name']}」 완료 "
                  f"({len(qs)}문항 수신)", flush=True)
            for q in qs:
                q.setdefault("source", f"{sec['name']} / {sec['start']}-{sec['end']}쪽")
            all_questions.extend(qs)
            prior_titles.extend(q.get("question", "")[:50] for q in qs)
    else:
        # 단원 정보 없음 → 균등 N등분 폴백
        per = max(total // FALLBACK_SECTIONS, 1)
        print(f"[QuizHigh] 균등 분할 폴백: {FALLBACK_SECTIONS}회 × {per}문항",
              flush=True)
        for i in range(1, FALLBACK_SECTIONS + 1):
            spec = _section_spec(
                section_name=f"영역 {i}",
                section_count=per,
                total=total,
                index=i,
                total_sections=FALLBACK_SECTIONS,
                prior_titles=prior_titles,
            )
            instruction = (
                prompt_tpl.replace("{{N}}", str(per))
                .replace("{{SECTION_SPEC}}", spec)
            )
            print(f"[QuizHigh] 분할 {i}/{FALLBACK_SECTIONS} 시작", flush=True)
            qs = await _call_section(rag, instruction)
            print(f"[QuizHigh] 분할 {i}/{FALLBACK_SECTIONS} 완료 "
                  f"({len(qs)}문항 수신)", flush=True)
            all_questions.extend(qs)
            prior_titles.extend(q.get("question", "")[:50] for q in qs)

    questions = all_questions[:total]
    print(f"[QuizHigh] 총 {len(questions)}문항 병합 완료", flush=True)

    files: list[Path] = []
    artifacts_dir = context.get("artifacts_dir")
    if artifacts_dir and questions:
        out_dir = Path(artifacts_dir) / META.key
        stamp = int(time.time())
        try:
            xlsx_path = write_table_xlsx(
                out_dir / f"quiz_high_{stamp}.xlsx",
                headers=["#", "문제", "보기1", "보기2", "보기3", "보기4",
                         "정답", "해설", "출처", "검수"],
                rows=_xlsx_rows(questions),
                sheet_name="고난도 30문제",
            )
            files.append(xlsx_path)
        except Exception as e:
            print(f"[quiz_high] xlsx 생성 실패: {e}", flush=True)
        try:
            docx_path = _write_docx(
                out_dir / f"quiz_high_{stamp}.docx",
                questions,
                deck_title=f"고난도 {len(questions)}문제",
            )
            files.append(docx_path)
        except Exception as e:
            print(f"[quiz_high] docx 생성 실패: {e}", flush=True)

    return ArtifactResult(
        key=META.key,
        title=META.title,
        markdown=_markdown(questions) or "_(문항 생성 실패)_",
        data={
            "questions": questions,
            "sections_detected": len(sections),
            "total_requested": total,
        },
        files=files,
    )
